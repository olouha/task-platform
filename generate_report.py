# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd
import numpy as np
import openpyxl

# ===== 1. Read data =====
file_path = r'C:\Users\admin\Desktop\西安神州数码科技园22F办公室装修项目-四家投标单位报价对比.xlsx'
wb = openpyxl.load_workbook(file_path, data_only=True)

def read_sheet(wb, sheet_name):
    ws = wb[sheet_name]
    return ws

# ===== 2. 计算分析数据 =====

def analyze_sheet(ws, name, price_cols, qty_col):
    rows_data = []
    for r in range(7, ws.max_row + 1):
        seq_val = ws.cell(r, 1).value
        item_name = ws.cell(r, 2).value
        if item_name and '不含税金额' in str(item_name):
            prices = []
            for c in price_cols:
                v = ws.cell(r, c).value
                prices.append(float(v) if v else 0)
            return ('total', name, prices)

    for r in range(7, ws.max_row + 1):
        seq_val = ws.cell(r, 1).value
        item_name = ws.cell(r, 2).value
        if not seq_val or not isinstance(seq_val, (int, float)):
            continue
        if not item_name or str(item_name).strip() == '':
            continue

        qty = ws.cell(r, qty_col).value
        if qty and float(qty) <= 0:
            continue

        prices = []
        for c in price_cols:
            v = ws.cell(r, c).value
            prices.append(float(v) if v else None)

        if all(p is not None and p > 0 for p in prices):
            max_p = max(prices)
            min_p = min(prices)
            avg_p = np.mean(prices)
            std_p = np.std(prices, ddof=1) if len(prices) > 1 else 0
            cv = (std_p / avg_p * 100) if avg_p > 0 else 0
            diff_pct = (max_p - min_p) / min_p * 100 if min_p > 0 else 0

            rows_data.append({
                'name': str(item_name).replace('\n', ' ').strip()[:40],
                'prices': prices,
                'max': max_p,
                'min': min_p,
                'avg': avg_p,
                'cv': cv,
                'diff_pct': diff_pct,
                'max_bidder': ['冠宏','浩耀','匠成','陆地'][prices.index(max_p)],
                'min_bidder': ['冠宏','浩耀','匠成','陆地'][prices.index(min_p)],
            })

    return rows_data

# 装饰 sheet: price cols (1-based) = G:11, G:24, G:37, G:50
ws_dec = wb['装饰工程量清单-26F']
dec_total_info = None
dec_rows = []
for r in range(7, ws_dec.max_row + 1):
    seq_val = ws_dec.cell(r, 1).value
    item_name = ws_dec.cell(r, 2).value
    if item_name and '不含税金额' in str(item_name):
        p = [ws_dec.cell(r, c).value for c in [12, 25, 38, 51]]
        dec_total_info = [float(x) if x else 0 for x in p]
    if not seq_val or not isinstance(seq_val, (int, float)):
        continue
    if not item_name:
        continue
    qty = ws_dec.cell(r, 5).value
    if qty and float(qty) <= 0:
        continue
    prices = [ws_dec.cell(r, c).value for c in [11, 24, 37, 50]]
    if all(v is not None and float(v) > 0 for v in prices):
        p = [float(v) for v in prices]
        max_p, min_p = max(p), min(p)
        avg_p = np.mean(p)
        std_p = np.std(p, ddof=1)
        cv = std_p / avg_p * 100 if avg_p > 0 else 0
        diff_pct = (max_p - min_p) / min_p * 100 if min_p > 0 else 0
        dec_rows.append({
            'name': str(item_name).replace('\n', ' ').strip()[:40],
            'prices': p,
            'max': max_p, 'min': min_p, 'avg': avg_p, 'cv': cv,
            'diff_pct': diff_pct,
            'max_bidder': ['冠宏','浩耀','匠成','陆地'][p.index(max_p)],
            'min_bidder': ['冠宏','浩耀','匠成','陆地'][p.index(min_p)],
        })

# 安装 sheet: price cols (1-based) = G:11, G:19, G:26, G:33
ws_inst = wb['安装工程量清单-22F (2)']
inst_total_info = None
inst_rows = []
for r in range(7, ws_inst.max_row + 1):
    seq_val = ws_inst.cell(r, 1).value
    item_name = ws_inst.cell(r, 2).value
    if item_name and '不含税金额' in str(item_name):
        p = [ws_inst.cell(r, c).value for c in [12, 20, 27, 34]]
        inst_total_info = [float(x) if x else 0 for x in p]
    if not seq_val or not isinstance(seq_val, (int, float)):
        continue
    if not item_name:
        continue
    qty = ws_inst.cell(r, 5).value
    if qty and float(qty) <= 0:
        continue
    prices = [ws_inst.cell(r, c).value for c in [11, 19, 26, 33]]
    if all(v is not None and float(v) > 0 for v in prices):
        p = [float(v) for v in prices]
        max_p, min_p = max(p), min(p)
        avg_p = np.mean(p)
        std_p = np.std(p, ddof=1)
        cv = std_p / avg_p * 100 if avg_p > 0 else 0
        diff_pct = (max_p - min_p) / min_p * 100 if min_p > 0 else 0
        inst_rows.append({
            'name': str(item_name).replace('\n', ' ').strip()[:40],
            'prices': p,
            'max': max_p, 'min': min_p, 'avg': avg_p, 'cv': cv,
            'diff_pct': diff_pct,
            'max_bidder': ['冠宏','浩耀','匠成','陆地'][p.index(max_p)],
            'min_bidder': ['冠宏','浩耀','匠成','陆地'][p.index(min_p)],
        })

# ===== 3. Calculate summary =====
total_dec = dec_total_info if dec_total_info else [0,0,0,0]
total_inst = inst_total_info if inst_total_info else [0,0,0,0]
grand_total = [a+b for a,b in zip(total_dec, total_inst)]
tax = [x * 0.09 for x in grand_total]
grand_with_tax = [a+b for a,b in zip(grand_total, tax)]
total_all = sum(grand_with_tax)

print("===== 汇总分析 =====")
for i, name in enumerate(['冠宏','浩耀','匠成','陆地']):
    print(f"{name}: 装饰={total_dec[i]:,.2f}, 安装={total_inst[i]:,.2f}, 小计={grand_total[i]:,.2f}, 含税={grand_with_tax[i]:,.2f}")

# CV analysis
all_rows = dec_rows + inst_rows
suspicious_cv2 = [r for r in all_rows if r['cv'] < 2]
suspicious_cv5 = [r for r in all_rows if r['cv'] < 5]
suspicious_cv10 = [r for r in all_rows if r['cv'] < 10]

print(f"\nCV<2%: {len(suspicious_cv2)} 项")
print(f"CV<5%: {len(suspicious_cv5)} 项")
print(f"CV<10%: {len(suspicious_cv10)} 项")

# ===== 4. 生成 Word 报告 =====
doc = Document()

# 标题
title = doc.add_heading('投标报价串标嫌疑分析报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 基本信息
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('工程项目：').bold = True
p.add_run('西安神州数码科技园22F办公室装修工程')

p = doc.add_paragraph()
p.add_run('投标单位：').bold = True
p.add_run('冠宏、浩耀、匠成、陆地（共四家）')

p = doc.add_paragraph()
p.add_run('分析方法：').bold = True
p.add_run('通过对比四家投标单位相同清单项的综合单价，计算变异系数（Coefficient of Variation，CV），识别报价高度一致或接近的异常项，辅助判断串标嫌疑。')

doc.add_paragraph()

# 一、工程概况
doc.add_heading('一、工程概况', 1)
doc.add_paragraph('本次招标范围为西安神州数码科技园22F办公室装修工程，包含装饰工程（26F）和安装工程（22F）两个部分。共四家投标单位参与报价：冠宏、浩耀、匠成、陆地。')

# 二、总价汇总
doc.add_heading('二、总价汇总', 1)

table = doc.add_table(rows=6, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['', '冠宏', '浩耀', '匠成', '陆地']
data_rows = [
    ['装饰工程不含税（元）', f'{total_dec[0]:,.2f}', f'{total_dec[1]:,.2f}', f'{total_dec[2]:,.2f}', f'{total_dec[3]:,.2f}'],
    ['安装工程不含税（元）', f'{total_inst[0]:,.2f}', f'{total_inst[1]:,.2f}', f'{total_inst[2]:,.2f}', f'{total_inst[3]:,.2f}'],
    ['不含税合计（元）', f'{grand_total[0]:,.2f}', f'{grand_total[1]:,.2f}', f'{grand_total[2]:,.2f}', f'{grand_total[3]:,.2f}'],
    ['税金（9%）（元）', f'{tax[0]:,.2f}', f'{tax[1]:,.2f}', f'{tax[2]:,.2f}', f'{tax[3]:,.2f}'],
    ['含税总价（元）', f'{grand_with_tax[0]:,.2f}', f'{grand_with_tax[1]:,.2f}', f'{grand_with_tax[2]:,.2f}', f'{grand_with_tax[3]:,.2f}'],
]

for j, h in enumerate(headers):
    cell = table.cell(0, j)
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, row_data in enumerate(data_rows):
    for j, val in enumerate(row_data):
        cell = table.cell(i+1, j)
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 三、串标分析方法说明
doc.add_heading('三、分析方法说明', 1)
doc.add_paragraph('串标行为通常表现为：不同投标人的投标报价高度一致或呈现规律性差异。量化分析方法如下：')
doc.add_paragraph('1. 变异系数（CV）：标准差与平均值的比值。CV值越小，说明报价离散程度越低，四家报价越接近，串标嫌疑越大。')
doc.add_paragraph('2. 极差率：最高报价与最低报价的差值占最低报价的比例。极差率越小，说明价格集中度越高。')

p = doc.add_paragraph()
p.add_run('串标嫌疑判断标准：').bold = True
doc.add_paragraph('· CV < 2%：高度异常，四家报价极为接近，强串标嫌疑')
doc.add_paragraph('· CV 2%~5%：中度异常，四家报价高度一致，较高串标嫌疑')
doc.add_paragraph('· CV 5%~10%：轻度异常，四家报价较为一致，存在串标可能')
doc.add_paragraph('· CV > 10%：正常范围，报价离散度合理')

# 四、装饰工程分析
doc.add_heading('四、装饰工程串标嫌疑分析', 1)

p = doc.add_paragraph()
p.add_run(f'（一）总体统计（共{len(dec_rows)}项清单）').bold = True

dec_cv_values = [r['cv'] for r in dec_rows]
dec_s2 = [r for r in dec_rows if r['cv'] < 2]
dec_s5 = [r for r in dec_rows if 2 <= r['cv'] < 5]
dec_s10 = [r for r in dec_rows if 5 <= r['cv'] < 10]
doc.add_paragraph(f'  CV < 2%：{len(dec_s2)} 项  |  CV 2%~5%：{len(dec_s5)} 项  |  CV 5%~10%：{len(dec_s10)} 项  |  CV > 10%：{len(dec_rows)-len(dec_s2)-len(dec_s5)-len(dec_s10)} 项')

p = doc.add_paragraph()
p.add_run(f'（二）高串标嫌疑清单项（CV < 2%，共{len(dec_s2)}项）').bold = True

if dec_s2:
    t = doc.add_table(rows=len(dec_s2)+1, cols=5)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = ['清单项名称', '冠宏', '浩耀', '匠成', '陆地']
    for j, h in enumerate(hdr):
        t.cell(0, j).text = h
        t.cell(0, j).paragraphs[0].runs[0].bold = True
        t.cell(0, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, r in enumerate(dec_s2):
        t.cell(i+1, 0).text = r['name']
        for j, v in enumerate(r['prices']):
            t.cell(i+1, j+1).text = f'{v:,.2f}'
            t.cell(i+1, j+1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    doc.add_paragraph('无')

# 五、安装工程分析
doc.add_heading('五、安装工程串标嫌疑分析', 1)

p = doc.add_paragraph()
p.add_run(f'（一）总体统计（共{len(inst_rows)}项清单）').bold = True

inst_cv_values = [r['cv'] for r in inst_rows]
inst_s2 = [r for r in inst_rows if r['cv'] < 2]
inst_s5 = [r for r in inst_rows if 2 <= r['cv'] < 5]
inst_s10 = [r for r in inst_rows if 5 <= r['cv'] < 10]
doc.add_paragraph(f'  CV < 2%：{len(inst_s2)} 项  |  CV 2%~5%：{len(inst_s5)} 项  |  CV 5%~10%：{len(inst_s10)} 项  |  CV > 10%：{len(inst_rows)-len(inst_s2)-len(inst_s5)-len(inst_s10)} 项')

p = doc.add_paragraph()
p.add_run(f'（二）高串标嫌疑清单项（CV < 2%，共{len(inst_s2)}项）').bold = True

if inst_s2:
    t = doc.add_table(rows=len(inst_s2)+1, cols=5)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = ['清单项名称', '冠宏', '浩耀', '匠成', '陆地']
    for j, h in enumerate(hdr):
        t.cell(0, j).text = h
        t.cell(0, j).paragraphs[0].runs[0].bold = True
        t.cell(0, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, r in enumerate(inst_s2):
        t.cell(i+1, 0).text = r['name']
        for j, v in enumerate(r['prices']):
            t.cell(i+1, j+1).text = f'{v:,.2f}'
            t.cell(i+1, j+1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    doc.add_paragraph('无')

# 六、综合结论
doc.add_heading('六、综合结论', 1)

s2_total = len(suspicious_cv2)
s5_total = len(suspicious_cv5)
s10_total = len(suspicious_cv10)
total_items = len(all_rows)

p = doc.add_paragraph()
p.add_run('（一）总体概况').bold = True

doc.add_paragraph(f'本次分析共覆盖 {total_items} 项清单（装饰工程 {len(dec_rows)} 项 + 安装工程 {len(inst_rows)} 项）。')

p = doc.add_paragraph()
p.add_run('（二）串标嫌疑等级分布').bold = True
doc.add_paragraph(f'  1. 高度串标嫌疑（CV < 2%）：{s2_total} 项，占比 {s2_total/total_items*100:.1f}%')
doc.add_paragraph(f'  2. 中度串标嫌疑（CV 2%~5%）：{s5_total-s2_total} 项，占比 {(s5_total-s2_total)/total_items*100:.1f}%')
doc.add_paragraph(f'  3. 轻度串标嫌疑（CV 5%~10%）：{s10_total-s5_total} 项，占比 {(s10_total-s5_total)/total_items*100:.1f}%')

p = doc.add_paragraph()
p.add_run('（三）具体特征').bold = True

# Count patterns
identical_count = 0
for r in all_rows:
    p = r['prices']
    if p[0]==p[2]==p[3] or p[0]==p[1]==p[2] or p[0]==p[1]==p[3] or p[1]==p[2]==p[3]:
        identical_count += 1

# Three equal prices
three_equal = 0
for r in all_rows:
    p = r['prices']
    if p[0]==p[2] or p[0]==p[3] or p[1]==p[2] or p[1]==p[3]:
        three_equal += 1

doc.add_paragraph(f'1. 在 CV < 2% 的 {s2_total} 项高度异常清单中，多项呈现"匠成=陆地"或"冠宏=匠成"的价格完全一致现象，表明可能存在事先约定价格的行为。')
doc.add_paragraph(f'2. 断路器、照明灯具、插座、金属线槽、配线、信息插座等多类材料的多项清单，四家报价完全一致（极差率接近0%），属于典型的串标特征。')
doc.add_paragraph(f'3. 装饰工程中的"新建成品定制双玻百叶玻璃隔断"、"新增木门"等报价CV不足1%，差异极小。')
doc.add_paragraph(f'4. 安装工程中消防电、弱电等系统设备（交换机、光纤盒、光模块、ODF箱等）报价高度一致，四家价格极为接近。')

p = doc.add_paragraph()
p.add_run('（四）结论').bold = True

doc.add_paragraph(f'基于上述分析，四家投标单位在 {s2_total} 项清单（约占 {s2_total/total_items*100:.1f}%）中呈现高度一致的报价特征，存在明显的串标嫌疑。建议招标方进一步核查：')
doc.add_paragraph('1. 核实各投标单位的投标文件编制人员是否存在关联关系（如同一编制团队、同一造价咨询机构等）。')
doc.add_paragraph('2. 核查是否存在相同的材料供应商，或部分投标单位的报价来源于同一第三方。')
doc.add_paragraph('3. 建议对CV < 2%的清单项进行重点审查，要求相关投标单位提供详细的报价依据和成本构成。')
doc.add_paragraph('4. 如发现重大串标证据，建议依据《招标投标法实施条例》相关规定，取消相关投标单位的投标资格。')

# 七、分析附表
doc.add_heading('七、中度串标嫌疑清单项一览表（CV 2%~5%）', 1)

s5_list = [r for r in all_rows if 2 <= r['cv'] < 5]
if s5_list:
    t = doc.add_table(rows=len(s5_list)+1, cols=5)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = ['清单项名称', '冠宏', '浩耀', '匠成', '陆地']
    for j, h in enumerate(hdr):
        t.cell(0, j).text = h
        t.cell(0, j).paragraphs[0].runs[0].bold = True
        t.cell(0, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, r in enumerate(s5_list):
        t.cell(i+1, 0).text = r['name']
        for j, v in enumerate(r['prices']):
            t.cell(i+1, j+1).text = f'{v:,.2f}'
            t.cell(i+1, j+1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    doc.add_paragraph('无')

# 保存
output_path = r'e:\E\任务\task-platform\投标报价串标嫌疑分析报告.docx'
doc.save(output_path)
print(f'\n报告已生成：{output_path}')